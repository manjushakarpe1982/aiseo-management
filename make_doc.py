#!/usr/bin/env python3
"""Generates AISEO_Claude_Data_Reference.docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "AISEO_Claude_Data_Reference.docx"

# ── Colours ────────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x2E, 0x75, 0xB6)
BLUE_LIGHT = RGBColor(0xD5, 0xE8, 0xF0)
GREY_LIGHT = RGBColor(0xF2, 0xF2, 0xF2)
GREEN_DARK = RGBColor(0x1F, 0x7A, 0x4F)
GREEN_LIGHT= RGBColor(0xE2, 0xEF, 0xDA)
ORANGE_DARK= RGBColor(0xC0, 0x55, 0x00)
ORANGE_LIGHT=RGBColor(0xFF, 0xF2, 0xCC)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
DARK_TEXT  = RGBColor(0x26, 0x26, 0x26)
CODE_BG    = RGBColor(0xF5, 0xF5, 0xF5)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'CCCCCC')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE_DARK
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = color or DARK_TEXT
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_TEXT
    return p


def body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def code_block(doc, text):
    """Monospace grey box for code/JSON."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)
    # shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    return p


def labelled_row(doc, label, value, label_bold=True):
    """Key: Value row in normal text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    r1.bold = label_bold
    r1.font.size = Pt(10)
    r2 = p.add_run(value)
    r2.font.size = Pt(10)
    return p


def add_table(doc, headers, rows, col_widths_cm=None, header_bg=BLUE_DARK, header_fg=WHITE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = header_fg
        set_cell_bg(cell, header_bg)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        bg = GREY_LIGHT if ridx % 2 == 0 else WHITE
        for cidx, val in enumerate(row_data):
            cell = row.cells[cidx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            set_cell_bg(cell, bg)

    # Column widths
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    set_cell_borders(table)
    doc.add_paragraph()
    return table


def banner(doc, text, bg: RGBColor, fg: RGBColor):
    """Full-width coloured label."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"  {text}  ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = fg
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    hex_bg = f"{bg[0]:02X}{bg[1]:02X}{bg[2]:02X}"
    hex_bg2 = f"{bg[0]:02X}{bg[1]:02X}{bg[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_bg2)
    pPr.append(shd)


# ══════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

# Default font
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# ── Title Page ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("AISEO — Claude API Data Reference")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = BLUE_DARK

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("What gets sent to Claude · How it is structured · What Claude returns")
run2.font.size = Pt(11)
run2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
run2.italic = True

doc.add_paragraph()

# ── Section 1: Overview ────────────────────────────────────────────────────
heading1(doc, "1. Overview")

body(doc,
     "The AISEO scanner makes two distinct Claude API calls per scan — one per URL "
     "tree cluster for cannibalization detection, and one per individual page for "
     "content improvement analysis. No other data is sent to Claude. All calls use "
     "the model claude-sonnet-4-20250514 with max_tokens=8192.")

add_table(doc,
    headers=["Call", "Trigger", "Frequency", "Model"],
    rows=[
        ["Cannibalization Analysis", "Once per URL tree cluster", "e.g. 1 call for 'silver-bars' tree with 80 pages", "claude-sonnet-4-20250514"],
        ["Content Improvement Analysis", "Once per individual page", "e.g. 250 calls for 250 pages", "claude-sonnet-4-20250514"],
    ],
    col_widths_cm=[4.5, 5, 5.5, 4.5]
)

body(doc,
     "A URL tree cluster is derived from the 2nd segment of the URL path. "
     "For example: /silver-bullion/silver-bars/1-oz → cluster = 'silver-bars'. "
     "All pages sharing the same cluster are analysed together for cannibalization.")

# ── Section 2: Scraped Fields ──────────────────────────────────────────────
heading1(doc, "2. Fields Scraped From Each Page")

body(doc,
     "Before any Claude call is made, Playwright scrapes each URL and stores the "
     "following fields in ClCode_ScrapedPages. These are the raw inputs available "
     "for Claude's analysis.")

add_table(doc,
    headers=["Field", "Source on Page", "Stored As", "Used in Claude Call"],
    rows=[
        ["PageURL",        "Final URL after Cloudflare redirect",         "NVARCHAR(500)",  "Both calls"],
        ["MetaTitle",      "document.title",                              "NVARCHAR(500)",  "Both calls"],
        ["MetaDescription","<meta name='description'>",                   "NVARCHAR(1000)", "Both calls"],
        ["H1",             "First <h1> innerText",                        "NVARCHAR(500)",  "Both calls"],
        ["H2s",            "All <h2> innerText values (JSON array)",       "NVARCHAR(MAX)",  "Both calls"],
        ["H3s",            "All <h3> innerText values (JSON array)",       "NVARCHAR(MAX)",  "Cannibalization only"],
        ["BodyContent",    "#seoContent div innerText (5,000 char limit)", "NVARCHAR(MAX)",  "Content Improvement only"],
        ["WordCount",      "len(BodyContent.split())",                    "INT",            "Content Improvement only"],
        ["CanonicalURL",   "<link rel='canonical'> href",                  "NVARCHAR(500)",  "Content Improvement only"],
        ["SchemaMarkup",   "All <script type='application/ld+json'>",      "NVARCHAR(MAX)",  "Content Improvement only"],
        ["InternalLinks",  "All same-domain <a href> (max 200)",           "NVARCHAR(MAX)",  "Not sent to Claude"],
        ["ImageAltTags",   "All <img alt> values",                        "NVARCHAR(MAX)",  "Not sent to Claude"],
    ],
    col_widths_cm=[4, 5.5, 3.5, 5]
)

body(doc,
     "Important: BodyContent is extracted exclusively from the #seoContent div "
     "(the dedicated SEO editorial block present on all pages of boldpreciousmetals.com). "
     "This avoids including product listing noise (prices, 'Add to Cart' buttons, "
     "header/footer/navigation). If #seoContent is absent, the scraper falls back to "
     "full-body extraction with product grids, nav, header, and footer removed.",
     italic=True, color=RGBColor(0x50, 0x50, 0x50))

# ── Section 3: Cannibalization Call ───────────────────────────────────────
heading1(doc, "3. Call 1 — Cannibalization Analysis")

labelled_row(doc, "Trigger",    "After all pages in a tree cluster are scraped")
labelled_row(doc, "Frequency",  "1 API call per tree cluster (e.g. 'silver-bars', '1-oz-silver-coins')")
labelled_row(doc, "Skipped if", "Tree cluster has fewer than 2 pages")
labelled_row(doc, "Output stored in", "ClCode_CannibalizationIssues")
doc.add_paragraph()

# 3a: Data structure
heading2(doc, "3a. Data Structure Sent (TREE_DATA)", color=BLUE_DARK)
body(doc,
     "The user message contains a JSON array of page objects. Each object "
     "has the following fields — note that BodyContent is truncated to a "
     "500-character excerpt for the cannibalization call (full body is not needed):")

code_block(doc, '''{TREE_DATA} = [
  {
    "url":             "https://www.boldpreciousmetals.com/silver-bullion/silver-bars/1-oz-silver-bars",
    "MetaTitle":       "1 oz Silver Bars | .999 Fine | BOLD Precious Metals",
    "MetaDescription": "Buy 1 oz silver bars online at BOLD...",
    "H1":              "Buy 1 oz Silver Bars",
    "H2s":             ["Shop by Refinery", "Why Buy Silver Bars?", ...],
    "H3s":             ["Low Premiums", "IRA Eligible", ...],
    "BodyContent":     "1 oz silver bars are the most popular..."  // first 500 chars only
  },
  {
    "url":             "https://www.boldpreciousmetals.com/silver-bullion/silver-bars/10-oz-silver-bars",
    "MetaTitle":       "10 oz Silver Bars | .999 Fine | BOLD Precious Metals",
    ...
  }
  // ... all pages in the cluster
]''')

# 3b: System Prompt
heading2(doc, "3b. System Prompt", color=BLUE_DARK)
body(doc, "The full system prompt sent to Claude for cannibalization analysis:")

code_block(doc, '''You are an expert SEO analyst specialising in ecommerce bullion websites.

You will be given a cluster of related product listing pages from the same
website tree. Your job is to identify keyword cannibalization issues where
multiple pages are competing for the same search terms.

STRICT RULES — you must follow these without exception:

1. ONLY use the data provided in the JSON. Do not infer, assume, or invent
   any page content, keywords, or intent not explicitly present in the fields.

2. A cannibalization issue ONLY qualifies if the SAME keyword phrase
   (verbatim or near-verbatim, allowing for word order variation) appears in
   the MetaTitle, H1, or MetaDescription of TWO OR MORE different pages.
   URL slug similarity alone is NOT sufficient to flag cannibalization.

3. For url1_current_content and url2_current_content — copy the value
   EXACTLY as it appears in the provided JSON.

4. Do NOT flag pages differentiated by weight/size specifier (e.g. "1 oz"
   vs "10 oz") as cannibalization — these serve distinct search intents.

5. If a page has a null, empty, or missing field, treat that field as absent.

6. Severity rules:
   - High:   identical or near-identical MetaTitle AND H1 across two pages
   - Medium: same primary keyword in MetaTitle or H1, differences elsewhere
   - Low:    overlap only in MetaDescription or H2s/H3s

7. If no clear evidence-based issue exists, return an empty array [].''')

# 3c: User Message Template
heading2(doc, "3c. User Message Template", color=BLUE_DARK)
body(doc,
     "The {TREE_NAME} and {TREE_DATA} placeholders are replaced at runtime "
     "with the actual cluster name and the JSON array shown above.")

code_block(doc, '''Analyse the following pages from the "{TREE_NAME}" cluster for keyword
cannibalization issues.

Pages data (JSON):
{TREE_DATA}

Each page has: URL, MetaTitle, MetaDescription, H1, H2s, H3s, BodyContent excerpt.

REMINDER: Only flag cannibalization where the same keyword phrase is explicitly
present in MetaTitle, H1, or MetaDescription of two or more pages.

For EACH cannibalization issue found, return a JSON array:
[
  {
    "cannibal_keyword":      "exact phrase found verbatim in both pages' fields",
    "severity":              "High|Medium|Low",
    "severity_reason":       "one sentence: which fields triggered this level",
    "url1":                  "first competing URL",
    "url1_field":            "MetaTitle|MetaDescription|H1|H2|H3",
    "url1_current_content":  "exact value copied from JSON",
    "url1_suggested_fix":    "specific rewrite removing overlap",
    "url2":                  "second competing URL",
    "url2_field":            "MetaTitle|MetaDescription|H1|H2|H3",
    "url2_current_content":  "exact value copied from JSON",
    "url2_suggested_fix":    "specific rewrite removing overlap",
    "overall_recommendation":"consolidate|differentiate|redirect|canonical",
    "reasoning":             "cite exact overlapping field values and SERP impact"
  }
]

Return ONLY the JSON array. No preamble. No markdown fences.
If no qualifying cannibalization is found, return: []''')

# 3d: Output fields
heading2(doc, "3d. Output — Fields Saved to ClCode_CannibalizationIssues", color=BLUE_DARK)

add_table(doc,
    headers=["JSON Field Returned", "DB Column", "Description"],
    rows=[
        ["cannibal_keyword",      "CannibalKeyword",       "The exact keyword phrase competing across pages"],
        ["severity",              "Severity",              "High / Medium / Low"],
        ["severity_reason",       "SeverityReason",        "One sentence explaining severity level"],
        ["url1",                  "URL1",                  "First competing page URL"],
        ["url1_field",            "URL1_FieldName",        "Which field (MetaTitle / H1 / etc.) has the overlap"],
        ["url1_current_content",  "URL1_CurrentContent",   "Exact current value of that field"],
        ["url1_suggested_fix",    "URL1_SuggestedFix",     "Rewrite suggestion for page 1"],
        ["url2",                  "URL2",                  "Second competing page URL"],
        ["url2_field",            "URL2_FieldName",        "Which field on page 2"],
        ["url2_current_content",  "URL2_CurrentContent",   "Exact current value on page 2"],
        ["url2_suggested_fix",    "URL2_SuggestedFix",     "Rewrite suggestion for page 2"],
        ["overall_recommendation","OverallRecommendation", "consolidate / differentiate / redirect / canonical"],
        ["reasoning",             "Reasoning",             "Full explanation citing exact field values"],
    ],
    col_widths_cm=[5, 5, 8.5]
)

# ── Section 4: Content Improvement Call ───────────────────────────────────
heading1(doc, "4. Call 2 — Content Improvement Analysis")

labelled_row(doc, "Trigger",    "After cannibalization analysis for each page's tree is complete")
labelled_row(doc, "Frequency",  "1 API call per individual page (e.g. 250 calls for 250 pages)")
labelled_row(doc, "Output stored in", "ClCode_ContentImprovements")
doc.add_paragraph()

# 4a: Data structure
heading2(doc, "4a. Data Structure Sent (PAGE_DATA)", color=GREEN_DARK)
body(doc,
     "A single JSON object is sent for each page. BodyContent is capped at "
     "5,000 characters (extracted from #seoContent — pure editorial copy, "
     "no product listing noise):")

code_block(doc, '''{PAGE_DATA} = {
  "url":             "https://www.boldpreciousmetals.com/silver-bullion/silver-coins/1-oz-silver-coins",
  "MetaTitle":       "Buy 1 oz Silver Coins | BOLD Precious Metals",
  "MetaDescription": "Shop 1 oz silver coins at BOLD. American Eagle, Maple Leaf...",
  "H1":              "Buy 1 oz Silver Coins",
  "H2s":             ["Why Choose BOLD?", "Shop by Mint", "FAQs"],
  "H3s":             ["Low Premiums", "IRA Eligible", "Fast Shipping"],
  "BodyContent":     "Secure your capital with 1 oz silver coins — the gold standard
                      for private investors seeking maximum liquidity, legal tender
                      status, and sovereign purity..."   // up to 5,000 chars from #seoContent
  "WordCount":       412,
  "CanonicalURL":    "https://www.boldpreciousmetals.com/silver-bullion/silver-coins/1-oz-silver-coins",
  "SchemaMarkup":    "[{\"@type\":\"Product\", \"name\":\"1 oz Silver Coin\", ...}]"
}''')

# 4b: System Prompt
heading2(doc, "4b. System Prompt", color=GREEN_DARK)

code_block(doc, '''You are an expert SEO content strategist specialising in precious metals
ecommerce. You will be given a product listing page with its current scraped
content. Suggest specific, field-level improvements to maximise organic
search rankings.

Note: this site has previously received Surfer SEO and Gemini suggestions,
so focus on improvements those tools typically miss — semantic depth, user
intent alignment, E-E-A-T signals, and keyword differentiation.

STRICT RULES — you must follow these without exception:

1. ONLY analyse the fields present in the JSON. Do not assume or invent
   content for null, empty string, or absent fields.

2. If a field is null or empty string — flag it as "Missing Field" with
   High priority.

3. Base every suggestion ONLY on the actual field values provided.

4. Field length thresholds (apply strictly, using character counts):
   - MetaTitle:       Too Short if < 50 chars | Too Long if > 60 chars
   - MetaDescription: Too Short if < 140 chars | Too Long if > 160 chars
   - H1:              Flag Missing Keyword only if the primary product
                      keyword is genuinely absent from the H1 value
   - BodyContent:     Flag Thin Content only if WordCount < 300

5. Do NOT flag an issue if the field already meets the threshold.

6. For current_content — copy the value EXACTLY as it appears in the JSON.

7. For suggested_content — provide a complete, ready-to-publish replacement.
   Do not use placeholders like [brand name] or [keyword].

8. Do not invent competitor data, ranking positions, or search volume figures.''')

# 4c: User Message Template
heading2(doc, "4c. User Message Template", color=GREEN_DARK)

code_block(doc, '''Analyse the following page for SEO content improvements.

Page data (JSON):
{PAGE_DATA}

REMINDER before you begin:
- Count MetaTitle characters from exact string. Flag Too Short only if < 50,
  Too Long only if > 60.
- Count MetaDescription characters. Flag only if < 140 or > 160.
- Flag Thin Content only if WordCount field value is under 300.
- Null or empty string fields = Missing Field — do not assume content.
- Copy current_content values exactly from JSON — do not alter them.
- Do not reference any information not present in the JSON above.

For EACH improvement needed, return a JSON array:
[
  {
    "field_name":          "MetaTitle|MetaDescription|H1|H2|BodyContent|SchemaMarkup",
    "current_content":     "exact value copied from JSON",
    "current_char_count":  0,
    "suggested_content":   "complete, ready-to-publish replacement — no placeholders",
    "suggested_char_count":0,
    "issue_type":          "Too Short|Too Long|Missing Keyword|Missing Field|Thin Content|
                            Keyword Stuffed|Poor Structure|No LSI|Duplicate",
    "reasoning":           "cite actual character count or keyword gap",
    "priority":            "High|Medium|Low",
    "impact_estimate":     "CTR Impact|Rankings|Featured Snippet|E-E-A-T|Crawlability"
  }
]

Return ONLY the JSON array. No preamble. No markdown fences.
If a field meets all thresholds and has no genuine issue, omit it.''')

# 4d: Output fields
heading2(doc, "4d. Output — Fields Saved to ClCode_ContentImprovements", color=GREEN_DARK)

add_table(doc,
    headers=["JSON Field Returned", "DB Column", "Description"],
    rows=[
        ["field_name",          "FieldName",          "Which field needs improvement (MetaTitle, H1, etc.)"],
        ["current_content",     "CurrentContent",     "Exact current value copied from the page"],
        ["current_char_count",  "CurrentCharCount",   "Character count of current content"],
        ["suggested_content",   "SuggestedContent",   "Complete ready-to-publish replacement"],
        ["suggested_char_count","SuggestedCharCount", "Character count of suggested replacement"],
        ["issue_type",          "IssueType",          "Too Short / Too Long / Missing Keyword / Missing Field / Thin Content / etc."],
        ["reasoning",           "Reasoning",          "Evidence-based explanation citing actual values"],
        ["priority",            "Priority",           "High / Medium / Low"],
        ["impact_estimate",     "ImpactEstimate",     "CTR Impact / Rankings / Featured Snippet / E-E-A-T / Crawlability"],
    ],
    col_widths_cm=[5, 5, 8.5]
)

# ── Section 5: What is NOT sent to Claude ──────────────────────────────────
heading1(doc, "5. What is NOT Sent to Claude")

add_table(doc,
    headers=["Field / Data", "Reason Not Sent"],
    rows=[
        ["InternalLinks",   "Not relevant to content quality or cannibalization analysis"],
        ["ImageAltTags",    "Not currently part of SEO analysis scope"],
        ["H4s, H5s, H6s",  "Too granular; H2/H3 provide sufficient heading structure"],
        ["Product prices",  "Stripped from BodyContent via #seoContent selector"],
        ["Navigation / header / footer text", "Stripped via #seoContent or DOM removal"],
        ["'Add to Cart' buttons / product grid", "Stripped via #seoContent selector"],
        ["Scrape timestamps", "Internal system metadata — not SEO-relevant"],
        ["ScanID / PromptID", "Internal DB keys — not content"],
        ["User credentials", "Never sent anywhere outside the SQL Server DB"],
    ],
    col_widths_cm=[7, 11.5]
)

# ── Section 6: Token Cost Estimate ────────────────────────────────────────
heading1(doc, "6. Approximate Token Cost Per Scan")

body(doc,
     "Estimated at ~4 characters per token. Actual costs depend on page "
     "content length and Claude's response length.")

add_table(doc,
    headers=["Call Type", "Typical Input Tokens", "Max Output Tokens", "Total per Call", "250-page scan estimate"],
    rows=[
        ["Cannibalization (per tree)",   "~3,000–5,000",  "8,192", "~11,000–13,000", "~15 trees × 12,000 = ~180,000"],
        ["Content Improvement (per page)","~2,500–4,000", "8,192", "~10,500–12,000", "250 pages × 11,000 = ~2,750,000"],
        ["TOTAL (full 250-page scan)",   "—",             "—",     "—",              "~2.9M tokens ≈ $4–8 USD"],
    ],
    col_widths_cm=[4.5, 4, 3.5, 3.5, 5]
)

body(doc,
     "The 3-second sleep between Claude calls (rate limit protection) means a "
     "250-page scan takes approximately 30–45 minutes to complete.",
     italic=True, color=RGBColor(0x55, 0x55, 0x55))

# ── Section 7: Data Flow ───────────────────────────────────────────────────
heading1(doc, "7. End-to-End Data Flow")

steps = [
    ("Phase 1", "Scan initialised",
     "Active prompts loaded from ClCode_Prompts. Source URLs fetched from AISEO_PageSEOInputs."),
    ("Phase 2", "Scraping (Playwright)",
     "Each URL is opened in a headless browser. #seoContent extracted as BodyContent. "
     "MetaTitle, MetaDescription, H1, H2s, H3s, CanonicalURL, SchemaMarkup saved to ClCode_ScrapedPages. "
     "Already-scraped URLs skipped (resume support)."),
    ("Phase 3", "Tree clustering",
     "Pages deduplicated (latest scrape per URL). Grouped by 2nd URL path segment → tree clusters."),
    ("Phase 4a", "Cannibalization API call",
     "For each tree with ≥2 pages: TREE_DATA JSON built → Claude called → "
     "results parsed → saved to ClCode_CannibalizationIssues."),
    ("Phase 4b", "Content Improvement API call",
     "For each page: PAGE_DATA JSON built → Claude called → "
     "results parsed → saved to ClCode_ContentImprovements."),
    ("Phase 5", "Completion + Report",
     "Scan status set to 'Completed'. Excel report optionally generated with 3 sheets: "
     "Scan Summary, Cannibalization Issues, Content Improvements."),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdrs = table.rows[0]
for i, h in enumerate(["Phase", "Step", "What Happens"]):
    hdrs.cells[i].text = ''
    r = hdrs.cells[i].paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = WHITE
    set_cell_bg(hdrs.cells[i], BLUE_DARK)

for phase, step, desc in steps:
    row = table.add_row()
    row.cells[0].text = phase
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
    row.cells[0].paragraphs[0].runs[0].bold = True
    set_cell_bg(row.cells[0], BLUE_LIGHT)

    row.cells[1].text = step
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
    row.cells[1].paragraphs[0].runs[0].bold = True

    row.cells[2].text = desc
    row.cells[2].paragraphs[0].runs[0].font.size = Pt(9.5)

for row in table.rows:
    row.cells[0].width = Cm(2)
    row.cells[1].width = Cm(4)
    row.cells[2].width = Cm(12.5)

set_cell_borders(table)
doc.add_paragraph()

# ── Footer note ───────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AISEO Management System  ·  Generated by Claude Code  ·  March 2026")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run.italic = True

# ── Save ──────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
